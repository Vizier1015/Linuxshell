import logging
import sys
import os
import datetime
import pymysql
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn



# log日志信息
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(filename)s[line:%(lineno)d] %(levelname)s %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')
logger = logging.getLogger(__name__)

# 读取的文件绝对路径
read_txt_path = os.path.join(os.getcwd(), 'search_table_list.txt')
sample_word_file_path = os.path.join(os.getcwd(), 'search_table_list.txt')

# 数据库连接信息 MYSQL
RDS_HOST = '*.*.*.*'
RDS_PORT = 3306
RDS_USER = ''
RDS_PASSWD = ''
RDS_DB = ''
RDS_TABLE = 'table_info_with_columns'


# odps的类
class MysqlClass(object):
    """用于获取表字段字典映射关系"""

    def __init__(self, **kwargs):
        # 多个参数 定义
        for k, w in kwargs.items():
            setattr(self, k, w)
        # 数据库连接信息
        self.conn = pymysql.connect(host=RDS_HOST, port=RDS_PORT, user=RDS_USER, passwd=RDS_PASSWD, db=RDS_DB,
                                    charset='utf8')
        # 创建游标链接
        self.cursor = self.conn.cursor()

    '''查询MYSQL数据库的的信息'''

    def search_mysql(self, execute_sql):
        # 查询的SQL语句
        result_list = []
        self.cursor.execute(execute_sql)
        data_list = self.cursor.fetchall()
        for value in data_list:
            result_list.append(value)
        return result_list


'''
********************************************
    读取txt中的数据内容，放入指定的list中
    要求：01.只支持一次数据一个主域的数据；02.子域需要先进行排序处理；03.数据以";"分隔
    project;table;table_name;class1;class2;
********************************************
'''


def load_ods_table(table_list_file):
    search_table_list = []
    with open(table_list_file, 'r', encoding='utf-8') as f:
        for line in f:
            if line is not None and line.strip('\n').strip() != "":
                search_table_list.append(line.strip('\n').strip())
    logging.info("读取文件成功，共有{}张表".format(len(search_table_list)))
    return search_table_list


'''
********************************************
    新增    table
    project;table;table_name;class1;class2;
********************************************
'''


def add_table_data_to_docx(document, get_queried_list):
    # 增加 表格
    table = document.add_table(rows=len(get_queried_list) + 1, cols=8)
    # 表格样式
    table.style = "Table Grid"
    table.autofit = True
    # 字体样式
    table.style.font.size = Pt(10.5)
    table.style.font.name = '黑体'
    # 新增第一行
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "字段"
    table.rows[0].cells[2].text = "字段名称"
    table.rows[0].cells[3].text = "字段类型"
    table.rows[0].cells[4].text = "字段格式"
    table.rows[0].cells[5].text = "代码集"
    table.rows[0].cells[6].text = "安全等级"
    table.rows[0].cells[7].text = "备注"

    # 遍历传入的数据
    for index, data_list in enumerate(get_queried_list):
        xh = str(index + 1) + ")"
        field = data_list[3]
        field_name = data_list[4]
        field_type = data_list[5]
        if field_type == "DOUBLE":
            presentation_format = "n10,6"
        elif field_type == "DATETIME":
            presentation_format = "d19"
        elif field == "dt":
            presentation_format = "c..8"
        else:
            presentation_format = "c..256"
        code_set = ""
        security_level = "L2"
        bz = ""
        table.rows[index + 1].cells[0].text = xh
        table.rows[index + 1].cells[1].text = field
        table.rows[index + 1].cells[2].text = field_name
        table.rows[index + 1].cells[3].text = field_type
        table.rows[index + 1].cells[4].text = presentation_format
        table.rows[index + 1].cells[5].text = code_set
        table.rows[index + 1].cells[6].text = security_level
        table.rows[index + 1].cells[7].text = bz


'''新增table'''


def add_table_data_to_docx_02(document, get_queried_list):
    # 增加 表格
    table = document.add_table(rows=len(get_queried_list) + 1, cols=4)
    # 表格样式
    table.style = "Table Grid"
    table.autofit = True
    # 字体样式
    table.style.font.size = Pt(10.5)
    table.style.font.name = '黑体'
    # 新增第一行
    table.rows[0].cells[0].text = "序号"
    # table.rows[0].cells[1].text = "三级类目"
    table.rows[0].cells[2].text = "数据表"
    # table.rows[0].cells[3].text = "安全级别"
    # 遍历传入的数据
    for index, list_value in enumerate(get_queried_list):
        xh = str(index + 1) + ")"
        value_arr = str(list_value).split(";")
        class_02 = value_arr[4]
        table_name = value_arr[2]
        table.rows[index + 1].cells[0].text = xh
        table.rows[index + 1].cells[1].text = class_02
        table.rows[index + 1].cells[2].text = table_name
        table.rows[index + 1].cells[3].text = "第二级"


'''获取主域、子域的数据'''


def get_class_list(list):
    class_02_dict = dict()
    for list_value in list:
        value_arr = str(list_value).split(";")
        class1 = value_arr[3]
        class2 = value_arr[2]
        class_02_dict[class1] = class2
    return class_02_dict


'''构造公共值'''


def comment_data_list():
    result_list = []
    data_tople_01 = ("", "", "", "dwd_zjid", "主键", "string")
    data_tople_02 = ("", "", "", "sjlyb", "数据来源表", "string")
    data_tople_03 = ("", "", "", "sjlyxt", "数据来源系统", "string")
    data_tople_04 = ("", "", "", "dwd_yxzt", "dwd层有效状态", "string")
    data_tople_05 = ("", "", "", "dwd_rksj", "dwd入库时间", "string")
    result_list.append(data_tople_01)
    result_list.append(data_tople_02)
    result_list.append(data_tople_03)
    result_list.append(data_tople_04)
    result_list.append(data_tople_05)
    return result_list


if __name__ == "__main__":
    # 开始执行时间
    starttime = datetime.datetime.now()
    logging.info("******************程序开始执行时间为{}*********************".format(starttime))
    # 01.读取数据，存入指定list
    search_table_list = load_ods_table(read_txt_path)
    # 创建文档对象
    document = Document()
    '''生成  （4概述）的word文档内容 '''
    if len(search_table_list) > 0:
        result_dict = get_class_list(search_table_list)
        document.add_heading("4概述", 1)
        '''4.1　文档生成'''
        # 计数器
        count = 0
        class_2 = ""
        for key, value in result_dict.items():
            class_2 = value
            if count == 0:
                document.add_heading("4.1　" + class_2 , 2)
            alphabet_str = chr(ord("a") + count)
            document.add_paragraph().add_run(alphabet_str + ")" + key)
            count += 1
        '''4.2　文档生成'''
        document.add_heading("4.2　" + class_2 + "域大数据的数据表", 2)
        document.add_paragraph().add_run("{}域大数据的数据表如下:".format(class_2))
        add_table_data_to_docx_02(document, search_table_list)

    '''生成  （5概述）的word文档内容 '''
    document.add_heading("5公共参数", 1)
    common_str = "以下数据字段为所有数据表的公共参数，每个数据表的字段均应包含公共参数。"
    document.add_paragraph().add_run(common_str)
    # 增加 表格
    writerlist = comment_data_list()
    add_table_data_to_docx(document, writerlist)

    '''
    **********************************************************
            生成公共参数后的word文档样式
    **********************************************************
    '''
    # 是否添加一级标题
    is_add_heading_01 = False
    class1 = ""
    # class2_01 = ""
    # 计数器
    count_01 = 6
    # count_02 = 1
    for index, list_value in enumerate(search_table_list):
        logging.info("正在处理第 {} 张表，生成文档较慢，请耐心等候！".format(index + 1))
        value_arr = str(list_value).split(";")
        table_project = value_arr[0]
        table = value_arr[1]
        table_name = value_arr[2]
        class1 = value_arr[3]
        class2 = value_arr[4]
        if class2_01 != class2:
            is_add_heading_01 = False
            class2_01 = class2
        # 获取的字段顺序混乱
        queried_sql = '''
                    select project_name,table_name,table_comment,column_name,column_comment,column_type
                    from {0}.{1} where  table_name = '{2}'
                '''.format(RDS_DB, RDS_TABLE, table)
        get_queried_list = MysqlClass().search_mysql(queried_sql)

        # 数据写入docx文档中
        if not is_add_heading_01:
            run = document.add_heading("", 1).add_run(str(count_01) + class2)
            run.font.name = u"微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            is_add_heading_01 = True
            count_01 += 1
            count_02 = 1
        # logging.info("表:{},状态：{};".format(table_name,is_add_heading_01))
        # 添加二级标题
        heading_name_2_1 = table_name
        document.add_heading(str(count_01 - 1) + "." + str(count_02) + " " + heading_name_2_1, 2)
        count_02 += 1
        # 追加表格数据到docx文档
        if get_queried_list is not None and len(get_queried_list) > 0:
            add_table_data_to_docx(document, get_queried_list)
    document.save("DWD层" + class1 + "域数据V1.1.docx")
    # 打印执行时间
    endtime = datetime.datetime.now()
    logging.info("程序开始执行时间为{},结束时间为{},运行时间为:{}".format(starttime, endtime, endtime - starttime))
